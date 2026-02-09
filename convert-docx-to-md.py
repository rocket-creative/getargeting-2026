#!/usr/bin/env python3
"""
Convert all .docx files in page-text directory to .md files
"""

import os
import subprocess
import sys
from pathlib import Path

# Directory containing the docx files
PAGE_TEXT_DIR = Path(__file__).parent / "page-text"
OUTPUT_DIR = PAGE_TEXT_DIR  # Output to same directory

def check_pandoc():
    """Check if pandoc is installed"""
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False

def convert_with_pandoc(docx_path: Path, md_path: Path):
    """Convert docx to md using pandoc"""
    try:
        subprocess.run([
            "pandoc",
            str(docx_path),
            "-f", "docx",
            "-t", "markdown",
            "-o", str(md_path),
            "--wrap=none"
        ], check=True, capture_output=True)
        return True
    except subprocess.CalledProcessError as e:
        print(f"  Error converting {docx_path.name}: {e.stderr.decode()}")
        return False

def convert_with_python(docx_path: Path, md_path: Path):
    """Fallback: Convert docx to md using python-docx"""
    try:
        from docx import Document
    except ImportError:
        print("Installing python-docx...")
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
        from docx import Document
    
    try:
        doc = Document(str(docx_path))
        md_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                md_content.append("")
                continue
            
            # Detect headings based on style
            style_name = para.style.name.lower() if para.style else ""
            
            if "heading 1" in style_name or "title" in style_name:
                md_content.append(f"# {text}")
            elif "heading 2" in style_name:
                md_content.append(f"## {text}")
            elif "heading 3" in style_name:
                md_content.append(f"### {text}")
            elif "heading 4" in style_name:
                md_content.append(f"#### {text}")
            elif "list" in style_name:
                md_content.append(f"- {text}")
            else:
                md_content.append(text)
        
        # Handle tables
        for table in doc.tables:
            md_content.append("")
            headers = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    headers = cells
                    md_content.append("| " + " | ".join(cells) + " |")
                    md_content.append("| " + " | ".join(["---"] * len(cells)) + " |")
                else:
                    md_content.append("| " + " | ".join(cells) + " |")
            md_content.append("")
        
        with open(md_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_content))
        
        return True
    except Exception as e:
        print(f"  Error converting {docx_path.name}: {e}")
        return False

def main():
    print("=" * 60)
    print("DOCX to Markdown Converter")
    print("=" * 60)
    
    # Check if page-text directory exists
    if not PAGE_TEXT_DIR.exists():
        print(f"Error: Directory not found: {PAGE_TEXT_DIR}")
        sys.exit(1)
    
    # Get all docx files
    docx_files = list(PAGE_TEXT_DIR.glob("*.docx"))
    
    if not docx_files:
        print("No .docx files found in page-text directory")
        sys.exit(0)
    
    print(f"\nFound {len(docx_files)} .docx files to convert\n")
    
    # Check for pandoc (preferred method)
    use_pandoc = check_pandoc()
    if use_pandoc:
        print("Using pandoc for conversion (recommended)\n")
    else:
        print("Pandoc not found. Using python-docx fallback.\n")
        print("For better results, install pandoc: brew install pandoc\n")
    
    # Convert each file
    success_count = 0
    fail_count = 0
    
    for docx_path in sorted(docx_files):
        md_path = docx_path.with_suffix(".md")
        print(f"Converting: {docx_path.name} -> {md_path.name}")
        
        if use_pandoc:
            success = convert_with_pandoc(docx_path, md_path)
        else:
            success = convert_with_python(docx_path, md_path)
        
        if success:
            success_count += 1
            print(f"  ✓ Success")
        else:
            fail_count += 1
            print(f"  ✗ Failed")
    
    print("\n" + "=" * 60)
    print(f"Conversion complete!")
    print(f"  Success: {success_count}")
    print(f"  Failed:  {fail_count}")
    print(f"\nMarkdown files saved to: {OUTPUT_DIR}")
    print("=" * 60)

if __name__ == "__main__":
    main()
